import pandas as pd
import numpy as np
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.model_selection import train_test_split
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.tree import DecisionTreeClassifier, plot_tree
from sklearn.metrics import accuracy_score, precision_recall_fscore_support, multilabel_confusion_matrix
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from docx.shared import Inches

df = pd.read_csv('train.csv')

agri_labels = [
    'barley', 'carcass', 'castor-oil', 'cocoa', 'coconut', 'coconut-oil', 'coffee', 
    'copra-cake', 'corn', 'cotton', 'cotton-oil', 'grain', 'groundnut', 'groundnut-oil', 
    'hog', 'l-cattle', 'lin-oil', 'livestock', 'meal-feed', 'oat', 'oilseed', 'orange', 
    'palm-oil', 'palmkernel', 'potato', 'rape-oil', 'rapeseed', 'rice', 'rubber', 'rye', 
    'sorghum', 'soy-meal', 'soy-oil', 'soybean', 'sugar', 'sun-meal', 'sun-oil', 'sunseed', 
    'tea', 'veg-oil', 'wheat'
]

df_target = df[['text'] + agri_labels].copy()

lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

def clean_text(text):
    if not isinstance(text, str): return ""
    text = text.lower()
    text = re.sub(r'[^a-z\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    tokens = text.split()
    cleaned_tokens = [lemmatizer.lemmatize(word) for word in tokens if word not in stop_words and len(word) > 1]
    return " ".join(cleaned_tokens)

df_target['cleaned_text'] = df_target['text'].apply(clean_text)

X = df_target['cleaned_text']
y = df_target[agri_labels].values

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

def evaluate_model(X_train_feat, X_test_feat, name):
    dt = DecisionTreeClassifier(random_state=42)
    dt.fit(X_train_feat, y_train)
    y_pred = dt.predict(X_test_feat)
    
    acc = accuracy_score(y_test, y_pred)
    p_mac, r_mac, f1_mac, _ = precision_recall_fscore_support(y_test, y_pred, average='macro', zero_division=0)
    p_mic, r_mic, f1_mic, _ = precision_recall_fscore_support(y_test, y_pred, average='micro', zero_division=0)
    
    return {
        'name': name,
        'acc': acc,
        'p_mac': p_mac, 'r_mac': r_mac, 'f1_mac': f1_mac,
        'p_mic': p_mic, 'r_mic': r_mic, 'f1_mic': f1_mic,
        'y_pred': y_pred
    }

print("Running models...")
# Eks-1: BoW
bow = CountVectorizer(max_features=1000)
X_train_bow = bow.fit_transform(X_train)
X_test_bow = bow.transform(X_test)
res_bow = evaluate_model(X_train_bow, X_test_bow, "Eks-1: DT + BoW")

# Eks-2: N-gram (Bigram)
ngram = CountVectorizer(ngram_range=(2, 2), max_features=1000)
X_train_ngram = ngram.fit_transform(X_train)
X_test_ngram = ngram.transform(X_test)
res_ngram = evaluate_model(X_train_ngram, X_test_ngram, "Eks-2: DT + N-gram")

# Eks-3: TF-IDF
tfidf = TfidfVectorizer(max_features=1000)
X_train_tfidf = tfidf.fit_transform(X_train)
X_test_tfidf = tfidf.transform(X_test)
res_tfidf = evaluate_model(X_train_tfidf, X_test_tfidf, "Eks-3: DT + TF-IDF")

results = [res_bow, res_ngram, res_tfidf]

print("Generating plots...")
labels = [r['name'] for r in results]
acc = [r['acc'] for r in results]
f1_mac = [r['f1_mac'] for r in results]
f1_mic = [r['f1_mic'] for r in results]

x = np.arange(len(labels))
width = 0.25

fig, ax = plt.subplots(figsize=(10, 6))
rects1 = ax.bar(x - width, acc, width, label='Accuracy')
rects2 = ax.bar(x, f1_mac, width, label='F1-Score (Macro)')
rects3 = ax.bar(x + width, f1_mic, width, label='F1-Score (Micro)')

ax.set_ylabel('Scores')
ax.set_title('Perbandingan Evaluasi Eksperimen 1-3')
ax.set_xticks(x)
ax.set_xticklabels(labels)
ax.legend()
plt.tight_layout()
plt.savefig('eval_metrics.png')
plt.close()

# Pohon keputusan penuh (tanpa batasan kedalaman)
dt_full = DecisionTreeClassifier(random_state=42)
dt_full.fit(X_train_tfidf, y_train)

# Visualisasi pohon keputusan (depth 5 agar tetap terbaca di dokumen)
plt.figure(figsize=(28, 14))
plot_tree(dt_full, feature_names=tfidf.get_feature_names_out(),
          max_depth=5, filled=True, rounded=True, fontsize=8,
          impurity=False, proportion=True)
plt.title(f'Pohon Keputusan (Decision Tree)\nKedalaman Penuh: {dt_full.get_depth()} | Jumlah Daun: {dt_full.get_n_leaves()}',
          fontsize=14, fontweight='bold')
plt.savefig('decision_tree_vis.png', dpi=150, bbox_inches='tight')
plt.close()

# Grafik Feature Importance (Top 20 kata paling berpengaruh)
importances = dt_full.feature_importances_
feat_names = tfidf.get_feature_names_out()
top_n = 20
top_idx = np.argsort(importances)[-top_n:]
top_importances = importances[top_idx]
top_feat_names = feat_names[top_idx]

plt.figure(figsize=(10, 8))
sns.barplot(x=top_importances, y=top_feat_names, palette='viridis')
plt.title(f'Top {top_n} Fitur Paling Berpengaruh dalam Decision Tree', fontsize=13, fontweight='bold')
plt.xlabel('Feature Importance Score')
plt.ylabel('Kata (Fitur)')
plt.tight_layout()
plt.savefig('feature_importance.png', dpi=150, bbox_inches='tight')
plt.close()

# Info statistik pohon
tree_info = {
    'depth': dt_full.get_depth(),
    'n_leaves': dt_full.get_n_leaves(),
    'n_features': dt_full.n_features_in_
}

top_5_labels = np.argsort(y.sum(axis=0))[-5:]
top_5_names = [agri_labels[i] for i in top_5_labels]

def plot_cm(y_pred, filename, title):
    mcm = multilabel_confusion_matrix(y_test, y_pred)
    fig, axes = plt.subplots(1, 5, figsize=(20, 4))
    for i, (label_idx, label_name) in enumerate(zip(top_5_labels, top_5_names)):
        sns.heatmap(mcm[label_idx], annot=True, fmt='d', cmap='Blues', ax=axes[i], cbar=False)
        axes[i].set_title(f'{label_name}')
        axes[i].set_xlabel('Predicted')
        axes[i].set_ylabel('Actual')
    plt.suptitle(title, y=1.05)
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()

plot_cm(res_bow['y_pred'], 'cm_bow.png', 'Confusion Matrix Top 5 Kategori - Eks 1 (BoW)')
plot_cm(res_ngram['y_pred'], 'cm_ngram.png', 'Confusion Matrix Top 5 Kategori - Eks 2 (N-gram)')
plot_cm(res_tfidf['y_pred'], 'cm_tfidf.png', 'Confusion Matrix Top 5 Kategori - Eks 3 (TF-IDF)')

print("Building docx...")
doc = docx.Document()
doc.add_heading('Tahap 5: Data Mining', 0)

doc.add_heading('1. Pohon Keputusan (Decision Tree) dan Grafik', level=2)
doc.add_paragraph('Pemodelan dilakukan menggunakan algoritma Decision Tree karena kemampuannya secara native untuk menangani multi-label text classification. Algoritma ini bekerja dengan cara memecah data berdasarkan aturan If-Else pada fitur kata-kata teks, sehingga proses pengambilan keputusan klasifikasi bisa kita telusuri secara transparan.')
doc.add_paragraph(f'Model Decision Tree yang dilatih menggunakan fitur TF-IDF (1.000 kata teratas) menghasilkan pohon keputusan dengan spesifikasi sebagai berikut:')

# Tabel info pohon
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_table.rows[0].cells[0].text = 'Parameter'
info_table.rows[0].cells[1].text = 'Nilai'
info_table.rows[1].cells[0].text = 'Kedalaman Pohon (Tree Depth)'
info_table.rows[1].cells[1].text = str(tree_info['depth'])
info_table.rows[2].cells[0].text = 'Jumlah Daun (Leaf Nodes)'
info_table.rows[2].cells[1].text = str(tree_info['n_leaves'])
info_table.rows[3].cells[0].text = 'Jumlah Fitur yang Digunakan'
info_table.rows[3].cells[1].text = str(tree_info['n_features'])

doc.add_paragraph('')
doc.add_paragraph('Visualisasi Pohon Keputusan:')
doc.add_paragraph('Berikut adalah visualisasi struktur pohon keputusan yang terbentuk. Karena pohon penuh sangat dalam dan lebar, tampilan di bawah ini menunjukkan 5 level teratas agar tetap dapat terbaca dengan jelas:')
doc.add_picture('decision_tree_vis.png', width=Inches(6.5))

doc.add_paragraph('')
doc.add_paragraph('Grafik Feature Importance:')
doc.add_paragraph('Grafik berikut menunjukkan 20 kata (fitur) yang paling berpengaruh dalam proses pengambilan keputusan oleh model Decision Tree. Semakin tinggi skor importance, semakin sering kata tersebut digunakan sebagai pemecah (splitter) di dalam pohon keputusan untuk membedakan antar kategori berita pertanian:')
doc.add_picture('feature_importance.png', width=Inches(5.5))

doc.add_heading('2. Skenario Eksperimen dan Visualisasi Evaluasi', level=2)
doc.add_paragraph('Model Decision Tree diuji menggunakan tiga skenario representasi teks yang berbeda:')
doc.add_paragraph('1. Eks-1: Decision Tree + Bag of Words (BoW)')
doc.add_paragraph('2. Eks-2: Decision Tree + N-gram (Bigram)')
doc.add_paragraph('3. Eks-3: Decision Tree + TF-IDF')
doc.add_paragraph('Berikut adalah visualisasi grafik batang yang membandingkan performa metrik Accuracy, F1-Score (Macro), dan F1-Score (Micro) untuk ketiga eksperimen:')
doc.add_picture('eval_metrics.png', width=Inches(6.0))
doc.add_paragraph('Tabel Lengkap Hasil Evaluasi:')
table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Eksperimen'
hdr_cells[1].text = 'Accuracy'
hdr_cells[2].text = 'Precision (Macro)'
hdr_cells[3].text = 'Recall (Macro)'
hdr_cells[4].text = 'F1-Score (Macro)'
hdr_cells[5].text = 'Precision (Micro)'
hdr_cells[6].text = 'Recall (Micro)'
hdr_cells[7].text = 'F1-Score (Micro)'

for r in results:
    row_cells = table.add_row().cells
    row_cells[0].text = r['name']
    row_cells[1].text = f"{r['acc']:.4f}"
    row_cells[2].text = f"{r['p_mac']:.4f}"
    row_cells[3].text = f"{r['r_mac']:.4f}"
    row_cells[4].text = f"{r['f1_mac']:.4f}"
    row_cells[5].text = f"{r['p_mic']:.4f}"
    row_cells[6].text = f"{r['r_mic']:.4f}"
    row_cells[7].text = f"{r['f1_mic']:.4f}"

doc.add_paragraph('\nBerdasarkan hasil tabel dan grafik di atas, dapat dilihat bahwa pendekatan menggunakan Term Frequency (TF-IDF) maupun Bag of Words memiliki akurasi dan metrik Micro-F1 yang stabil untuk menangani mayoritas kategori dalam dataset berita ini. Sedangkan N-gram (Bigram) menghasilkan performa yang paling rendah akibat tingginya tingkat sparsity dimensi teks.')

doc.add_heading('3. Confusion Matrix: Indikator Kinerja Model DT', level=2)
doc.add_paragraph('Karena data kita berformat multi-label, maka dihasilkan sekumpulan confusion matrix untuk setiap label kategori pertanian. Sebagai representasi, di bawah ini divisualisasikan indikator False Positive, True Positive, dsb. untuk 5 kategori mayoritas tertinggi di dataset berita Reuters ini secara berturut-turut untuk ketiga eksperimen.')

doc.add_heading('Eks-1: Confusion Matrix (DT + BoW)', level=3)
doc.add_picture('cm_bow.png', width=Inches(6.0))

doc.add_heading('Eks-2: Confusion Matrix (DT + N-gram)', level=3)
doc.add_picture('cm_ngram.png', width=Inches(6.0))

doc.add_heading('Eks-3: Confusion Matrix (DT + TF-IDF)', level=3)
doc.add_picture('cm_tfidf.png', width=Inches(6.0))

doc.add_paragraph('Dengan melihat ketiga hasil Confusion Matrix tersebut, metode BoW dan TF-IDF memberikan perbandingan matriks True Positives (prediksi tepat ke target aktual) yang paling optimal dibandingkan Bigram. Prediksi label pertanian pada dataset Reuters ini sebagian besar cukup kuat dikenali dengan pembobotan unigram tunggal secara langsung (seperti kata grain, corn, wheat) dibandingkan pemecahan multikata secara beruntun.')

doc.save('Tahap_5_Data_Mining_v2.docx')
print("Document saved as Tahap_5_Data_Mining_v2.docx")
