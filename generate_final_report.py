import pandas as pd
import numpy as np
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.model_selection import train_test_split
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.tree import DecisionTreeClassifier, export_text
from sklearn.metrics import (accuracy_score, precision_recall_fscore_support, 
                             multilabel_confusion_matrix, classification_report)
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

# ============================================================
# LOAD & PREPROCESS
# ============================================================
print("[1/6] Loading data...")
df = pd.read_csv('train.csv')

agri_labels = [
    'barley', 'carcass', 'castor-oil', 'cocoa', 'coconut', 'coconut-oil', 'coffee', 
    'copra-cake', 'corn', 'cotton', 'cotton-oil', 'grain', 'groundnut', 'groundnut-oil', 
    'hog', 'l-cattle', 'lin-oil', 'livestock', 'meal-feed', 'oat', 'oilseed', 'orange', 
    'palm-oil', 'palmkernel', 'potato', 'rape-oil', 'rapeseed', 'rice', 'rubber', 'rye', 
    'sorghum', 'soy-meal', 'soy-oil', 'soybean', 'sugar', 'sun-meal', 'sun-oil', 'sunseed', 
    'tea', 'veg-oil', 'wheat'
]

df_target = df[['text'] + agri_labels].copy()

lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

def clean_text(text):
    if not isinstance(text, str): return ""
    text = text.lower()
    text = re.sub(r'[^a-z\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    tokens = text.split()
    cleaned_tokens = [lemmatizer.lemmatize(word) for word in tokens if word not in stop_words and len(word) > 1]
    return " ".join(cleaned_tokens)

df_target['cleaned_text'] = df_target['text'].apply(clean_text)

X = df_target['cleaned_text']
y = df_target[agri_labels].values

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# ============================================================
# TRAIN 3 MODELS
# ============================================================
print("[2/6] Training 3 Decision Tree models...")

# Eks-1: BoW
bow = CountVectorizer(max_features=1000)
X_train_bow = bow.fit_transform(X_train)
X_test_bow = bow.transform(X_test)
dt_bow = DecisionTreeClassifier(random_state=42, max_depth=5)
dt_bow.fit(X_train_bow, y_train)
y_pred_bow = dt_bow.predict(X_test_bow)

# Eks-2: N-gram
ngram = CountVectorizer(ngram_range=(2, 2), max_features=1000)
X_train_ngram = ngram.fit_transform(X_train)
X_test_ngram = ngram.transform(X_test)
dt_ngram = DecisionTreeClassifier(random_state=42, max_depth=5)
dt_ngram.fit(X_train_ngram, y_train)
y_pred_ngram = dt_ngram.predict(X_test_ngram)

# Eks-3: TF-IDF
tfidf = TfidfVectorizer(max_features=1000)
X_train_tfidf = tfidf.fit_transform(X_train)
X_test_tfidf = tfidf.transform(X_test)
dt_tfidf = DecisionTreeClassifier(random_state=42, max_depth=5)
dt_tfidf.fit(X_train_tfidf, y_train)
y_pred_tfidf = dt_tfidf.predict(X_test_tfidf)

# ============================================================
# EXPORT TREE RULES (TEXT)
# ============================================================
print("[3/6] Exporting tree rules...")
rules_bow = export_text(dt_bow, feature_names=list(bow.get_feature_names_out()), max_depth=5)
rules_ngram = export_text(dt_ngram, feature_names=list(ngram.get_feature_names_out()), max_depth=5)
rules_tfidf = export_text(dt_tfidf, feature_names=list(tfidf.get_feature_names_out()), max_depth=5)

# ============================================================
# COMPUTE METRICS
# ============================================================
print("[4/6] Computing metrics...")

def get_metrics(y_true, y_pred):
    acc = accuracy_score(y_true, y_pred)
    p_mac, r_mac, f1_mac, _ = precision_recall_fscore_support(y_true, y_pred, average='macro', zero_division=0)
    p_mic, r_mic, f1_mic, _ = precision_recall_fscore_support(y_true, y_pred, average='micro', zero_division=0)
    return {
        'acc': acc,
        'p_mac': p_mac, 'r_mac': r_mac, 'f1_mac': f1_mac,
        'p_mic': p_mic, 'r_mic': r_mic, 'f1_mic': f1_mic
    }

m_bow = get_metrics(y_test, y_pred_bow)
m_ngram = get_metrics(y_test, y_pred_ngram)
m_tfidf = get_metrics(y_test, y_pred_tfidf)

# ============================================================
# GENERATE ALL PLOTS
# ============================================================
print("[5/6] Generating plots...")

# Top 5 label untuk confusion matrix
top_5_idx = np.argsort(y.sum(axis=0))[-5:]
top_5_names = [agri_labels[i] for i in top_5_idx]

# --- Feature Importance x3 ---
def plot_feature_importance(model, feat_names, title, filename):
    importances = model.feature_importances_
    top_n = 15
    top_idx = np.argsort(importances)[-top_n:]
    plt.figure(figsize=(9, 6))
    sns.barplot(x=importances[top_idx], y=np.array(feat_names)[top_idx], hue=np.array(feat_names)[top_idx], palette='viridis', legend=False)
    plt.title(title, fontsize=12, fontweight='bold')
    plt.xlabel('Importance Score')
    plt.ylabel('Fitur')
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight')
    plt.close()

plot_feature_importance(dt_bow, bow.get_feature_names_out(), 'Top 15 Feature Importance - Eks-1 (DT + BoW)', 'fi_bow.png')
plot_feature_importance(dt_ngram, ngram.get_feature_names_out(), 'Top 15 Feature Importance - Eks-2 (DT + N-gram)', 'fi_ngram.png')
plot_feature_importance(dt_tfidf, tfidf.get_feature_names_out(), 'Top 15 Feature Importance - Eks-3 (DT + TF-IDF)', 'fi_tfidf.png')

# --- Confusion Matrix x3 ---
def plot_cm(y_pred, filename, title):
    mcm = multilabel_confusion_matrix(y_test, y_pred)
    fig, axes = plt.subplots(1, 5, figsize=(22, 4))
    for i, (label_idx, label_name) in enumerate(zip(top_5_idx, top_5_names)):
        sns.heatmap(mcm[label_idx], annot=True, fmt='d', cmap='Blues', ax=axes[i], cbar=False,
                    xticklabels=['Neg (0)', 'Pos (1)'], yticklabels=['Neg (0)', 'Pos (1)'])
        axes[i].set_title(f'{label_name}', fontweight='bold')
        axes[i].set_xlabel('Predicted')
        axes[i].set_ylabel('Actual')
    plt.suptitle(title, fontsize=13, fontweight='bold', y=1.02)
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight')
    plt.close()

plot_cm(y_pred_bow, 'cm_bow.png', 'Confusion Matrix - Eks-1 (DT + BoW)')
plot_cm(y_pred_ngram, 'cm_ngram.png', 'Confusion Matrix - Eks-2 (DT + N-gram)')
plot_cm(y_pred_tfidf, 'cm_tfidf.png', 'Confusion Matrix - Eks-3 (DT + TF-IDF)')

# --- Perbandingan Evaluasi (Bar Chart) ---
fig, axes = plt.subplots(1, 3, figsize=(18, 5))
metrics_names = ['Accuracy', 'Precision\n(Macro)', 'Recall\n(Macro)', 'F1-Score\n(Macro)', 'Precision\n(Micro)', 'Recall\n(Micro)', 'F1-Score\n(Micro)']

for idx, (m, title, color) in enumerate([
    (m_bow, 'Eks-1: DT + BoW', '#3498db'),
    (m_ngram, 'Eks-2: DT + N-gram', '#e74c3c'),
    (m_tfidf, 'Eks-3: DT + TF-IDF', '#2ecc71')
]):
    vals = [m['acc'], m['p_mac'], m['r_mac'], m['f1_mac'], m['p_mic'], m['r_mic'], m['f1_mic']]
    axes[idx].barh(metrics_names, vals, color=color, edgecolor='gray')
    axes[idx].set_xlim(0, 1)
    axes[idx].set_title(title, fontweight='bold')
    for j, v in enumerate(vals):
        axes[idx].text(v + 0.01, j, f'{v:.4f}', va='center', fontsize=9)
plt.suptitle('Perbandingan Metrik Evaluasi Ketiga Eksperimen', fontsize=14, fontweight='bold')
plt.tight_layout()
plt.savefig('eval_comparison.png', dpi=150, bbox_inches='tight')
plt.close()

# ============================================================
# BUILD DOCX
# ============================================================
print("[6/6] Building docx...")
doc = docx.Document()

# ---- TITLE ----
doc.add_heading('Tahap 5 (Poin: 25): Data Mining', level=1)
doc.add_paragraph('Algoritma data mining yang digunakan: Decision Tree (Pohon Keputusan)')
doc.add_paragraph('Skenario eksperimen: 3 percobaan dengan representasi fitur teks yang berbeda')
doc.add_paragraph('Data dibagi menjadi 80% data latih dan 20% data uji menggunakan train_test_split (random_state=42).')

# ============ SECTION 1 ============
doc.add_heading('1. Pohon Keputusan dan Grafik', level=2)
doc.add_paragraph('Decision Tree merupakan algoritma klasifikasi yang bekerja dengan membagi data secara rekursif berdasarkan fitur-fitur tertentu menggunakan aturan if-else. Dalam proyek ini, model Decision Tree dilatih untuk melakukan multi-label classification pada 41 kategori pertanian. Berikut adalah pohon keputusan dan grafik feature importance untuk masing-masing eksperimen.')

# --- Eks-1: BoW ---
doc.add_heading('Eks-1: Pohon Keputusan (DT + Bag of Words)', level=3)
doc.add_paragraph(f'Spesifikasi pohon: Kedalaman = {dt_bow.get_depth()}, Jumlah Daun = {dt_bow.get_n_leaves()}, Jumlah Fitur = {dt_bow.n_features_in_}')
doc.add_paragraph('Aturan percabangan pohon keputusan:')
p = doc.add_paragraph()
run = p.add_run(rules_bow)
run.font.size = Pt(8)
run.font.name = 'Consolas'

doc.add_paragraph('Grafik Feature Importance (15 fitur terpenting):')
doc.add_picture('fi_bow.png', width=Inches(5.5))

# --- Eks-2: N-gram ---
doc.add_heading('Eks-2: Pohon Keputusan (DT + N-gram / Bigram)', level=3)
doc.add_paragraph(f'Spesifikasi pohon: Kedalaman = {dt_ngram.get_depth()}, Jumlah Daun = {dt_ngram.get_n_leaves()}, Jumlah Fitur = {dt_ngram.n_features_in_}')
doc.add_paragraph('Aturan percabangan pohon keputusan:')
p = doc.add_paragraph()
run = p.add_run(rules_ngram)
run.font.size = Pt(8)
run.font.name = 'Consolas'

doc.add_paragraph('Grafik Feature Importance (15 fitur terpenting):')
doc.add_picture('fi_ngram.png', width=Inches(5.5))

# --- Eks-3: TF-IDF ---
doc.add_heading('Eks-3: Pohon Keputusan (DT + TF-IDF)', level=3)
doc.add_paragraph(f'Spesifikasi pohon: Kedalaman = {dt_tfidf.get_depth()}, Jumlah Daun = {dt_tfidf.get_n_leaves()}, Jumlah Fitur = {dt_tfidf.n_features_in_}')
doc.add_paragraph('Aturan percabangan pohon keputusan:')
p = doc.add_paragraph()
run = p.add_run(rules_tfidf)
run.font.size = Pt(8)
run.font.name = 'Consolas'

doc.add_paragraph('Grafik Feature Importance (15 fitur terpenting):')
doc.add_picture('fi_tfidf.png', width=Inches(5.5))

# ============ SECTION 2 ============
doc.add_heading('2. Evaluasi Eksperimen', level=2)
doc.add_paragraph('Ketiga model Decision Tree dievaluasi menggunakan data uji (20% dari total dataset). Metrik yang diukur meliputi Accuracy, Precision, Recall, dan F1-Score baik secara Macro (rata-rata per kelas) maupun Micro (rata-rata global).')

doc.add_paragraph('Tabel Perbandingan Hasil Evaluasi:')
table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
headers = ['Eksperimen', 'Accuracy', 'Precision (Macro)', 'Recall (Macro)', 'F1 (Macro)', 'Precision (Micro)', 'Recall (Micro)', 'F1 (Micro)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

for name, m in [('Eks-1: DT + BoW', m_bow), ('Eks-2: DT + N-gram', m_ngram), ('Eks-3: DT + TF-IDF', m_tfidf)]:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = f"{m['acc']:.4f}"
    row[2].text = f"{m['p_mac']:.4f}"
    row[3].text = f"{m['r_mac']:.4f}"
    row[4].text = f"{m['f1_mac']:.4f}"
    row[5].text = f"{m['p_mic']:.4f}"
    row[6].text = f"{m['r_mic']:.4f}"
    row[7].text = f"{m['f1_mic']:.4f}"

doc.add_paragraph('')
doc.add_paragraph('Visualisasi Perbandingan Metrik:')
doc.add_picture('eval_comparison.png', width=Inches(6.5))

doc.add_paragraph('')
doc.add_paragraph('Analisis Hasil:')
doc.add_paragraph('Dari tabel dan grafik di atas, terlihat bahwa Eksperimen 1 (BoW) dan Eksperimen 3 (TF-IDF) memiliki performa yang hampir setara dan jauh lebih baik dibandingkan Eksperimen 2 (N-gram). Hal ini terjadi karena representasi N-gram (Bigram) menghasilkan fitur yang sangat sparse (jarang muncul), sehingga Decision Tree kesulitan menemukan pola yang konsisten untuk memisahkan kategori. Sementara itu, BoW dan TF-IDF yang menggunakan kata tunggal (unigram) lebih efektif karena kata-kata kunci pertanian seperti "wheat", "corn", "sugar" sudah cukup kuat untuk membedakan kategori.')

# ============ SECTION 3 ============
doc.add_heading('3. Confusion Matrix: Indikator Kinerja Model DT', level=2)
doc.add_paragraph('Confusion Matrix digunakan untuk mengukur kinerja model secara detail pada setiap kategori. Karena dataset ini bersifat multi-label (41 kategori), maka confusion matrix ditampilkan untuk 5 kategori pertanian dengan frekuensi tertinggi. Setiap matriks menunjukkan jumlah True Negative (TN), False Positive (FP), False Negative (FN), dan True Positive (TP).')
doc.add_paragraph(f'5 kategori yang ditampilkan: {", ".join(top_5_names)}')

# --- CM Eks-1 ---
doc.add_heading('Eks-1: Confusion Matrix (DT + BoW)', level=3)
doc.add_picture('cm_bow.png', width=Inches(6.5))

doc.add_paragraph('Tabel Detail Metrik Per Kategori (Eks-1):')
t1 = doc.add_table(rows=1, cols=5)
t1.style = 'Table Grid'
for i, h in enumerate(['Kategori', 'Accuracy', 'Precision', 'Recall', 'F1-Score']):
    t1.rows[0].cells[i].text = h
mcm_bow = multilabel_confusion_matrix(y_test, y_pred_bow)
for label_idx, label_name in zip(top_5_idx, top_5_names):
    tn, fp, fn, tp = mcm_bow[label_idx].ravel()
    total = tn + fp + fn + tp
    acc_l = (tp + tn) / total if total > 0 else 0
    prec_l = tp / (tp + fp) if (tp + fp) > 0 else 0
    rec_l = tp / (tp + fn) if (tp + fn) > 0 else 0
    f1_l = 2 * prec_l * rec_l / (prec_l + rec_l) if (prec_l + rec_l) > 0 else 0
    row = t1.add_row().cells
    row[0].text = label_name
    row[1].text = f"{acc_l:.4f}"
    row[2].text = f"{prec_l:.4f}"
    row[3].text = f"{rec_l:.4f}"
    row[4].text = f"{f1_l:.4f}"

# --- CM Eks-2 ---
doc.add_heading('Eks-2: Confusion Matrix (DT + N-gram)', level=3)
doc.add_picture('cm_ngram.png', width=Inches(6.5))

doc.add_paragraph('Tabel Detail Metrik Per Kategori (Eks-2):')
t2 = doc.add_table(rows=1, cols=5)
t2.style = 'Table Grid'
for i, h in enumerate(['Kategori', 'Accuracy', 'Precision', 'Recall', 'F1-Score']):
    t2.rows[0].cells[i].text = h
mcm_ngram = multilabel_confusion_matrix(y_test, y_pred_ngram)
for label_idx, label_name in zip(top_5_idx, top_5_names):
    tn, fp, fn, tp = mcm_ngram[label_idx].ravel()
    total = tn + fp + fn + tp
    acc_l = (tp + tn) / total if total > 0 else 0
    prec_l = tp / (tp + fp) if (tp + fp) > 0 else 0
    rec_l = tp / (tp + fn) if (tp + fn) > 0 else 0
    f1_l = 2 * prec_l * rec_l / (prec_l + rec_l) if (prec_l + rec_l) > 0 else 0
    row = t2.add_row().cells
    row[0].text = label_name
    row[1].text = f"{acc_l:.4f}"
    row[2].text = f"{prec_l:.4f}"
    row[3].text = f"{rec_l:.4f}"
    row[4].text = f"{f1_l:.4f}"

# --- CM Eks-3 ---
doc.add_heading('Eks-3: Confusion Matrix (DT + TF-IDF)', level=3)
doc.add_picture('cm_tfidf.png', width=Inches(6.5))

doc.add_paragraph('Tabel Detail Metrik Per Kategori (Eks-3):')
t3 = doc.add_table(rows=1, cols=5)
t3.style = 'Table Grid'
for i, h in enumerate(['Kategori', 'Accuracy', 'Precision', 'Recall', 'F1-Score']):
    t3.rows[0].cells[i].text = h
mcm_tfidf = multilabel_confusion_matrix(y_test, y_pred_tfidf)
for label_idx, label_name in zip(top_5_idx, top_5_names):
    tn, fp, fn, tp = mcm_tfidf[label_idx].ravel()
    total = tn + fp + fn + tp
    acc_l = (tp + tn) / total if total > 0 else 0
    prec_l = tp / (tp + fp) if (tp + fp) > 0 else 0
    rec_l = tp / (tp + fn) if (tp + fn) > 0 else 0
    f1_l = 2 * prec_l * rec_l / (prec_l + rec_l) if (prec_l + rec_l) > 0 else 0
    row = t3.add_row().cells
    row[0].text = label_name
    row[1].text = f"{acc_l:.4f}"
    row[2].text = f"{prec_l:.4f}"
    row[3].text = f"{rec_l:.4f}"
    row[4].text = f"{f1_l:.4f}"

doc.add_paragraph('')
doc.add_paragraph('Kesimpulan: Berdasarkan seluruh hasil evaluasi, model Decision Tree dengan representasi Bag of Words (Eks-1) dan TF-IDF (Eks-3) menunjukkan kinerja yang paling baik dalam mengklasifikasikan artikel berita ke dalam kategori pertanian. Sedangkan N-gram (Eks-2) kurang efektif karena fitur bigram terlalu spesifik dan jarang muncul secara konsisten di data uji. Untuk tahap selanjutnya, representasi TF-IDF direkomendasikan sebagai fitur utama karena memberikan keseimbangan terbaik antara precision dan recall.')

doc.save('Tahap_5_Data_Mining_FINAL.docx')
print("\nDocument saved: Tahap_5_Data_Mining_FINAL.docx")
